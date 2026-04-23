"""One-off patch: rewrite the "Prompt structure" section of the existing
baseline summary docx with verbatim, end-to-end prompts (no `{parent}`
templates, no elided demos). Preserves the rest of the document — including
the comparison table the user edited by hand.

Run from Project/code/:
    python _patch_prompts_section.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

HERE         = Path(__file__).resolve().parent
PROJECT_ROOT = HERE.parent
DOCX_PATH    = PROJECT_ROOT / "reports" / "baseline_summary_2026-04-23.docx"

sys.path.insert(0, str(HERE))
from baselines.celebrity_api.prompts import build_child_query
from baselines.llama_inference.cot_prompts import (
    build_fewshot_cot_reverse,
    build_hint_reverse,
    build_zeroshot_cot_reverse,
)

# Verbatim test pair. Mary Lee Pfeiffer is row 44 of the public CSV
# (parent_type=mother, can_reverse=False) — i.e. a paradigmatic
# reverse-direction test case. Not in the few-shot CoT demo set
# (Obama / Musk / Hemsworth), so the few-shot prompt isn't trivialised.
EXAMPLE_PARENT = "Mary Lee Pfeiffer"
EXAMPLE_CHILD  = "Tom Cruise"


def insert_text_before(target, text, *, bold=False, italic=False, size=None):
    p = target.insert_paragraph_before("")
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def insert_chat_before(target, messages):
    p = target.insert_paragraph_before("")
    p.paragraph_format.left_indent = Inches(0.25)
    for i, m in enumerate(messages):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(f"[{m['role']}] {m['content']}")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    return p


def main() -> None:
    doc = Document(str(DOCX_PATH))

    # Locate boundaries by exact text.
    header_p   = None   # "Prompt structure (each user turn shown verbatim; reverse direction)"
    template_p = None   # the big templates block (the next paragraph after header)
    scoring_p  = None   # "Scoring function" — insertion point

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if header_p is None and t.startswith("Prompt structure"):
            header_p = p
        elif header_p is not None and template_p is None and scoring_p is None:
            # First paragraph after the header (the templates block) —
            # but only if it isn't already "Scoring function".
            if t == "Scoring function":
                scoring_p = p
            else:
                template_p = p
        elif scoring_p is None and t == "Scoring function":
            scoring_p = p

    if header_p is None or scoring_p is None:
        raise RuntimeError(
            f"Could not locate boundaries: header_p={header_p}, "
            f"template_p={template_p}, scoring_p={scoring_p}"
        )

    # Delete OLD prompts paragraphs (header + template block).
    for old in (header_p, template_p):
        if old is not None:
            old._element.getparent().remove(old._element)

    # Insert NEW content immediately before "Scoring function".
    insert_text_before(
        scoring_p,
        "Prompt structure (exact prompts the model sees)",
        bold=True,
    )
    insert_text_before(
        scoring_p,
        f"Each block below is the verbatim chat-format input for the test pair "
        f"“{EXAMPLE_PARENT}” → “{EXAMPLE_CHILD}” "
        f"(asking each model to name a child of {EXAMPLE_PARENT}). System + all "
        f"few-shot turns + the test user turn are shown. The gold answer the "
        f"scorer is looking for is “{EXAMPLE_CHILD}”.",
        italic=True, size=10,
    )

    insert_text_before(scoring_p, "Direct (baseline)", bold=True)
    insert_text_before(
        scoring_p,
        "Original Berglund 2023 prompt: 3 mixed-direction few-shot demos "
        "(including one “I don't know.” to teach the abstain pattern), "
        "then the reverse-direction test query.",
        italic=True, size=10,
    )
    insert_chat_before(scoring_p, build_child_query(EXAMPLE_PARENT))

    insert_text_before(scoring_p, "Hint (condition 2)", bold=True)
    insert_text_before(
        scoring_p,
        "No demos. Single user turn that includes a hint asserting the parent "
        "has a famous child (no factual content beyond that).",
        italic=True, size=10,
    )
    insert_chat_before(scoring_p, build_hint_reverse(EXAMPLE_PARENT))

    insert_text_before(scoring_p, "Zero-shot CoT (condition 3b)", bold=True)
    insert_text_before(
        scoring_p,
        "No demos. Single user turn appending “Let's think step by step "
        "before answering.”",
        italic=True, size=10,
    )
    insert_chat_before(scoring_p, build_zeroshot_cot_reverse(EXAMPLE_PARENT))

    insert_text_before(scoring_p, "Few-shot CoT (condition 4)", bold=True)
    insert_text_before(
        scoring_p,
        "3 worked CoT demos (Obama, Musk, Hemsworth), each ending with "
        "“Answer: <Name>.” so the model learns the answer-span format. "
        "Then the test query in the same shape. Test pairs whose names overlap "
        "with any demo name are filtered before dispatch (the Hemsworth family "
        "is the only overlap in our 200-pair slice, accounting for 2 leaked "
        "pairs).",
        italic=True, size=10,
    )
    insert_chat_before(scoring_p, build_fewshot_cot_reverse(EXAMPLE_PARENT))

    doc.save(str(DOCX_PATH))
    print(f"Patched {DOCX_PATH}")


if __name__ == "__main__":
    main()
